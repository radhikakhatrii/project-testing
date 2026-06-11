"""
Deep Agent — a LangChain `deepagents` harness with planning, a file workspace,
and progressively-disclosed Skills (Excel, Word, PDF, data analysis).

UI: wide two-column layout — a chat column and a live side panel (Files with
preview, Plan/todos, and clickable Skills).

Design notes:
- Files are not auto-surfaced to the model, so each turn we inject the current
  workspace file list (the agent uses real names instead of guessing).
- Skill activation is otherwise model-driven and unreliable, so we deterministically
  route to the matching skill and load its SKILL.md into the turn context.
- Binary .xlsx/.docx are produced by in-process custom tools (no sandbox/code exec).
"""

# =============================================================================
# IMPORTS
# =============================================================================

import os
import re
import json
import shutil
import uuid
import zipfile
import tempfile

import streamlit as st
import nest_asyncio

nest_asyncio.apply()

import pandas as pd
from docx import Document
from pypdf import PdfReader

from langchain_openai import ChatOpenAI
from langchain_core.tools import tool

from deepagents import create_deep_agent
from deepagents.backends import FilesystemBackend

from utils.styles import apply_deep_agent_css


def _dbg(*parts):
    """Diagnostic line -> terminal (if attached) and tmp/deep_agent_debug.log.
    Guarded: under Streamlit on Windows sys.stdout may raise OSError(22)."""
    line = "[DeepAgent] " + " ".join(str(p) for p in parts)
    try:
        print(line, flush=True)
    except Exception:
        pass
    try:
        os.makedirs("tmp", exist_ok=True)
        with open(os.path.join("tmp", "deep_agent_debug.log"), "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except Exception:
        pass


# =============================================================================
# PAGE SETUP
# =============================================================================

st.set_page_config(page_title="Deep Agent Skill Architecture", page_icon="🤖", layout="centered",
                   initial_sidebar_state="expanded")

apply_deep_agent_css()

st.markdown('<div class="da-title">Deep Agent Skill Architecture</div>', unsafe_allow_html=True)
st.markdown('<div class="da-sub">Plans its work, operates over a file workspace, and loads Skills on demand.</div>', unsafe_allow_html=True)


# =============================================================================
# SESSION STATE
# =============================================================================

_DEFAULTS = {
    "da_session_id": uuid.uuid4().hex[:12],
    "da_agent": None,
    "da_messages": [],
    "da_todos": [],
    "da_last_trace": [],
    "da_last_skill_reads": [],
    "da_seen_files": [],
    "da_uploaded": [],
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# =============================================================================
# OPENAI KEY
# =============================================================================

openai_key = st.session_state.get("openai_key", "")
if not openai_key:
    st.markdown(
        '<div class="warn-banner">No OpenAI API key found. Go back to the Home page and save your key first.</div>',
        unsafe_allow_html=True)
    if st.button("Go to Home"):
        st.switch_page("Home.py")
    st.stop()


# =============================================================================
# WORKSPACE + SKILLS
# =============================================================================

SKILLS_SRC = os.path.abspath("skills")
WORKSPACE = os.path.abspath(os.path.join("tmp", "deep_agent", st.session_state.da_session_id))
SKILLS_DST = os.path.join(WORKSPACE, "skills")

os.makedirs(WORKSPACE, exist_ok=True)
# Mirror skills/ into the session workspace on every run, so any skill the user
# adds or edits is picked up automatically (new folders copied in, edited SKILL.md
# refreshed). Just drop a folder with a SKILL.md into skills/ and it appears.
if os.path.isdir(SKILLS_SRC):
    shutil.copytree(SKILLS_SRC, SKILLS_DST, dirs_exist_ok=True)


def _ws_path(filename: str) -> str:
    return os.path.join(WORKSPACE, os.path.basename(str(filename).strip().lstrip("/\\")))


def _workspace_files():
    try:
        return [n for n in os.listdir(WORKSPACE)
                if os.path.isfile(os.path.join(WORKSPACE, n)) and not n.startswith(".")]
    except FileNotFoundError:
        return []


def _not_found(name: str) -> str:
    avail = _workspace_files()
    return (f"Error: '{os.path.basename(name)}' is not in the workspace. "
            f"Available files: {avail if avail else 'none'}. Use one of these exact names.")


def _parse_skill(md_path: str):
    """Return (name, description, body) parsed from a SKILL.md file."""
    name = os.path.basename(os.path.dirname(md_path))
    desc, body = "", ""
    try:
        text = open(md_path, encoding="utf-8").read()
    except Exception:
        return name, desc, body
    body = text
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) == 3:
            fm, body = parts[1], parts[2].strip()
            for line in fm.splitlines():
                if line.startswith("name:"):
                    name = line.split(":", 1)[1].strip()
                elif line.startswith("description:"):
                    desc = line.split(":", 1)[1].strip()
    return name, desc, body


def _skill_details():
    out = []
    if not os.path.isdir(SKILLS_DST):
        return out
    for folder in sorted(os.listdir(SKILLS_DST)):
        md = os.path.join(SKILLS_DST, folder, "SKILL.md")
        if os.path.isfile(md):
            out.append(_parse_skill(md))
    return out


def _load_skill_body(folder: str) -> str:
    return _parse_skill(os.path.join(SKILLS_DST, folder, "SKILL.md"))[2]


# Common words to ignore when matching a request against a skill's description.
_SKILL_STOP = {
    "the", "a", "an", "of", "to", "for", "and", "or", "with", "this", "that", "when",
    "use", "used", "using", "you", "your", "are", "its", "from", "into", "user", "asks",
    "ask", "want", "wants", "need", "needs", "rather", "than", "just", "via", "them",
    "their", "make", "help", "provide", "based", "files", "file", "data", "content",
    "contents", "existing", "real", "new", "get",
}


def _kw(text: str) -> set:
    """Significant lowercase keywords from a string (for skill matching)."""
    return {w for w in re.findall(r"[a-z0-9]+", (text or "").lower())
            if len(w) >= 3 and w not in _SKILL_STOP}


def _skill_catalog():
    """[(folder, keyword_set)] for every skill on disk, derived from each skill's
    own name + description — so matching works for ANY skill, no names hard-coded."""
    cat = []
    if os.path.isdir(SKILLS_DST):
        for folder in sorted(os.listdir(SKILLS_DST)):
            md = os.path.join(SKILLS_DST, folder, "SKILL.md")
            if os.path.isfile(md):
                name, desc, _ = _parse_skill(md)
                cat.append((folder, _kw(name) | _kw(desc)))
    return cat


def _match_skills(text: str, files) -> list:
    """Pick the skill(s) most relevant to this turn, scored against each skill's own
    description + the uploaded file types. Fully generic: drop a new folder with a
    SKILL.md into skills/ (with a clear `description:`) and it routes automatically."""
    msg = _kw(text)
    exts = {os.path.splitext(f)[1].lower().lstrip(".") for f in files}
    scored = []
    for folder, skill_kw in _skill_catalog():
        score = len(msg & skill_kw) + (3 if (exts & skill_kw) else 0)
        if score > 0:
            scored.append((score, folder))
    if not scored:
        return []
    scored.sort(reverse=True)
    top = scored[0][0]
    # keep the best-matching skill(s) (allow a near-tie); cap at 2 to keep context lean
    return [f for s, f in scored if s >= max(1, top - 1)][:2]


def _sanitize_xlsx(path: str) -> str:
    """Strip XML comments from an .xlsx (some tool/LLM-generated files embed
    comments in workbook.xml that break openpyxl)."""
    fd, out = tempfile.mkstemp(suffix=".xlsx")
    os.close(fd)
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith((".xml", ".rels")):
                data = re.sub(rb"<!--.*?-->", b"", data, flags=re.S)
            zout.writestr(item, data)
    return out


def _read_excel(path: str, **kwargs):
    try:
        return pd.read_excel(path, **kwargs)
    except Exception:
        return pd.read_excel(_sanitize_xlsx(path), **kwargs)


# =============================================================================
# WORKSPACE FILE TOOLS
# =============================================================================

@tool
def create_excel(filename: str, sheets_json: str) -> str:
    """Create an Excel (.xlsx) workbook in the workspace.

    Args:
        filename: target file name ending in .xlsx (e.g. "budget.xlsx").
        sheets_json: JSON string mapping each sheet name to {"columns": [...],
            "rows": [[...], ...]}. Example:
            {"Sales": {"columns": ["Region", "Revenue"], "rows": [["North", 100]]}}
    """
    try:
        if not filename.lower().endswith(".xlsx"):
            filename += ".xlsx"
        data = json.loads(sheets_json)
        if not isinstance(data, dict) or not data:
            return "Error: sheets_json must be a non-empty JSON object {sheet: {columns, rows}}."
        path = _ws_path(filename)
        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            for sheet, spec in data.items():
                cols = spec.get("columns") if isinstance(spec, dict) else None
                rows = spec.get("rows", []) if isinstance(spec, dict) else spec
                pd.DataFrame(rows, columns=cols).to_excel(
                    writer, sheet_name=str(sheet)[:31], index=False)
        return f"Created '{os.path.basename(path)}' with sheets: {list(data.keys())}."
    except Exception as e:
        return f"Error creating Excel file: {type(e).__name__}: {e}"


@tool
def read_spreadsheet(filename: str) -> str:
    """Read a spreadsheet (.xlsx or .csv) and return a preview (shape + first rows)."""
    try:
        path = _ws_path(filename)
        if not os.path.isfile(path):
            return _not_found(filename)
        out = []
        if path.lower().endswith(".csv"):
            df = pd.read_csv(path)
            out.append(f"CSV '{os.path.basename(path)}' — {df.shape[0]} rows x {df.shape[1]} cols")
            out.append(df.head(20).to_string(index=False))
        else:
            for name, df in _read_excel(path, sheet_name=None).items():
                out.append(f"Sheet '{name}' — {df.shape[0]} rows x {df.shape[1]} cols")
                out.append(df.head(20).to_string(index=False))
        return "\n\n".join(out)
    except Exception as e:
        return f"Error reading spreadsheet: {type(e).__name__}: {e}"


@tool
def analyze_data(filename: str) -> str:
    """Analyze a tabular dataset (.csv or .xlsx): shape, dtypes, missing values,
    descriptive statistics, and a sample of rows."""
    try:
        path = _ws_path(filename)
        if not os.path.isfile(path):
            return _not_found(filename)
        df = pd.read_csv(path) if path.lower().endswith(".csv") else _read_excel(path)
        return "\n\n".join([
            f"Dataset '{os.path.basename(path)}' — {df.shape[0]} rows x {df.shape[1]} columns",
            "Column types:\n" + df.dtypes.astype(str).to_string(),
            "Missing values per column:\n" + df.isna().sum().to_string(),
            "Summary statistics:\n" + df.describe(include="all").transpose().to_string(),
            "Sample rows:\n" + df.head(5).to_string(index=False),
        ])
    except Exception as e:
        return f"Error analyzing data: {type(e).__name__}: {e}"


@tool
def create_word(filename: str, markdown_text: str) -> str:
    """Create a Word (.docx) document from simple Markdown ('# '/'## '/'### '
    headings, '- '/'* ' bullets, '1. ' numbered lists, blank-line paragraphs)."""
    try:
        if not filename.lower().endswith(".docx"):
            filename += ".docx"
        path = _ws_path(filename)
        doc = Document()
        for raw in markdown_text.split("\n"):
            line = raw.strip()
            if not line:
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif len(line) > 2 and line[0].isdigit() and line[1:3] == ". ":
                doc.add_paragraph(line[3:], style="List Number")
            else:
                doc.add_paragraph(line)
        doc.save(path)
        return f"Created '{os.path.basename(path)}'."
    except Exception as e:
        return f"Error creating Word document: {type(e).__name__}: {e}"


@tool
def read_document(filename: str) -> str:
    """Read the text of a Word (.docx) or plain-text (.txt) file."""
    try:
        path = _ws_path(filename)
        if not os.path.isfile(path):
            return _not_found(filename)
        if path.lower().endswith(".docx"):
            text = "\n".join(p.text for p in Document(path).paragraphs)
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        return (text[:20000] + "\n...[truncated]") if len(text) > 20000 else (text or "(empty)")
    except Exception as e:
        return f"Error reading document: {type(e).__name__}: {e}"


@tool
def read_pdf(filename: str) -> str:
    """Extract text from a PDF in the workspace, labeled per page."""
    try:
        path = _ws_path(filename)
        if not os.path.isfile(path):
            return _not_found(filename)
        chunks, total = [], 0
        for i, page in enumerate(PdfReader(path).pages, start=1):
            t = (page.extract_text() or "").strip()
            chunks.append(f"--- Page {i} ---\n{t}")
            total += len(t)
            if total > 20000:
                chunks.append("...[truncated remaining pages]")
                break
        if total == 0:
            return "The PDF has no extractable text (it may be scanned/image-only)."
        return "\n\n".join(chunks).strip()
    except Exception as e:
        return f"Error reading PDF: {type(e).__name__}: {e}"


@tool
def create_pdf(filename: str, markdown_text: str) -> str:
    """Create a PDF (.pdf) in the workspace from Markdown — headings, **bold**/
    *italic*, links, ordered/unordered/nested lists, and tables. Rendered via
    Markdown -> HTML -> fpdf2 (pure-Python, no system dependencies)."""
    try:
        import markdown as _md
        from fpdf import FPDF

        if not filename.lower().endswith(".pdf"):
            filename += ".pdf"
        path = _ws_path(filename)

        def _norm(s: str) -> str:
            # fpdf2 core fonts are latin-1 only — normalize common unicode, drop the rest
            for k, v in {"—": "-", "–": "-", "‘": "'", "’": "'",
                         "“": '"', "”": '"', "…": "...", "•": "-"}.items():
                s = s.replace(k, v)
            return s.encode("latin-1", "replace").decode("latin-1")

        html = _md.markdown(_norm(markdown_text),
                            extensions=["tables", "fenced_code", "sane_lists"])
        # fpdf2's HTML renderer styles <b>/<i>; map Markdown's <strong>/<em>
        html = (html.replace("<strong>", "<b>").replace("</strong>", "</b>")
                    .replace("<em>", "<i>").replace("</em>", "</i>"))

        def _build(use_html):
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Helvetica", size=11)
            if use_html:
                pdf.write_html(html)
            else:  # fallback: still produce a readable PDF if HTML rendering fails
                pdf.multi_cell(0, 6, _norm(markdown_text), new_x="LMARGIN", new_y="NEXT")
            return pdf

        try:
            pdf = _build(True)
        except Exception:
            pdf = _build(False)
        pdf.output(path)
        return f"Created '{os.path.basename(path)}'."
    except Exception as e:
        return f"Error creating PDF: {type(e).__name__}: {e}"


WORKSPACE_TOOLS = [create_excel, read_spreadsheet, analyze_data,
                   create_word, read_document, read_pdf, create_pdf]


# =============================================================================
# DEEP AGENT
# =============================================================================

SYSTEM_PROMPT = (
    "You are a Deep Agent. You plan your work, operate over a file workspace, and use Skills.\n\n"
    "Rules:\n"
    "1. ALWAYS begin by calling `write_todos` to lay out a short step-by-step plan, and keep it "
    "updated (mark steps in_progress/completed) as you go. Do this for every request.\n"
    "2. The latest user message lists the files currently in the workspace. Always use those "
    "EXACT filenames — never invent one. If the user says 'the file/sheet/document' and exactly "
    "one matches, use it. If unsure what exists, call `ls`.\n"
    "3. When the task context includes a 'Skill loaded' block, FOLLOW those instructions.\n"
    "4. Use the dedicated tools to read or create files: create_excel + read_spreadsheet/analyze_data "
    "for spreadsheets, create_word/read_document for Word & text, and read_pdf/create_pdf for PDFs. "
    "You can CREATE .xlsx, .docx and .pdf. If the user asks for a format you have NO tool for (e.g. "
    "PowerPoint/.pptx), do NOT fake it or write it with write_file — clearly tell them you can't "
    "create that format and offer the closest supported one (.docx, .xlsx, or .pdf). Pass plain "
    "filenames; created files are saved to the workspace and offered as downloads.\n"
    "5. CARRY THE TASK TO COMPLETION IN THIS SAME TURN. Never end your turn by only describing a "
    "plan or saying things like 'next I will…' / 'I will now create…'. If the request needs a file, "
    "you MUST actually call the create tool and confirm the file was created BEFORE you reply with "
    "text. Do NOT ask the user for permission or confirmation to proceed — just do it.\n"
    "6. Answer EXACTLY what was asked — produce the specific deliverable requested (the file, the "
    "analysis, the answer), nothing missing and nothing extra.\n"
    "7. FINISH as soon as the requested output exists: end with one short confirmation that names "
    "the file(s) you created and directly answers the request. Then stop."
)


def build_agent():
    _dbg("building agent | workspace:", WORKSPACE)
    _dbg("skills dir:", os.path.isdir(SKILLS_DST),
         "| contents:", os.listdir(SKILLS_DST) if os.path.isdir(SKILLS_DST) else "MISSING")
    llm = ChatOpenAI(model="gpt-4.1", temperature=0, api_key=openai_key)
    backend = FilesystemBackend(root_dir=WORKSPACE, virtual_mode=True)
    try:
        from deepagents.middleware.skills import _list_skills
        _dbg("skills discovered:", [s.get("name") for s in _list_skills(backend, "/skills")])
    except Exception as e:
        _dbg("skill discovery check FAILED:", repr(e))
    return create_deep_agent(
        model=llm, tools=WORKSPACE_TOOLS, backend=backend,
        skills=["/skills"], system_prompt=SYSTEM_PROMPT,
    )


# Rebuild the agent whenever the set of skills changes, so newly added/removed
# skills are reflected in the agent itself (not just the UI and routing).
_skills_sig = ",".join(
    f for f in (sorted(os.listdir(SKILLS_DST)) if os.path.isdir(SKILLS_DST) else [])
    if os.path.isfile(os.path.join(SKILLS_DST, f, "SKILL.md"))
)
if st.session_state.da_agent is None or st.session_state.get("da_skills_sig") != _skills_sig:
    with st.spinner("Initializing deep agent and loading skills..."):
        try:
            st.session_state.da_agent = build_agent()
            st.session_state.da_skills_sig = _skills_sig
        except Exception as e:
            st.error(f"Failed to initialize the deep agent: {type(e).__name__}: {e}")
            st.stop()

_n_skills = len(_skill_details())
st.markdown(
    f'<span class="da-status"><span class="da-dot"></span>Ready · planning · file workspace · '
    f'{_n_skills} skill{"s" if _n_skills != 1 else ""}</span>',
    unsafe_allow_html=True)


# =============================================================================
# SIDEBAR (controls)
# =============================================================================

with st.sidebar:
    st.markdown("**Deep Agent Skill Architecture**")
    st.caption("Plans, works over files, loads skills on demand.")
    st.divider()
    if st.button("Clear chat", use_container_width=True):
        st.session_state.da_messages = []
        st.session_state.da_todos = []
        st.session_state.da_last_trace = []
        st.session_state.da_last_skill_reads = []
        st.rerun()
    if st.button("Clear workspace", use_container_width=True):
        shutil.rmtree(WORKSPACE, ignore_errors=True)
        for k in ("da_messages", "da_todos", "da_last_trace", "da_last_skill_reads",
                  "da_seen_files", "da_uploaded"):
            st.session_state[k] = []
        st.session_state.da_agent = None
        st.rerun()
    if st.button("Home", use_container_width=True):
        st.switch_page("Home.py")


# =============================================================================
# HELPERS: history rendering + streamed-event labels
# =============================================================================

def _msg_parts(m):
    if isinstance(m, dict):
        return m.get("role", "assistant"), m.get("content", "")
    role = {"human": "user", "ai": "assistant", "system": "assistant",
            "tool": "assistant"}.get(getattr(m, "type", ""), "assistant")
    content = getattr(m, "content", "")
    if isinstance(content, list):
        content = "".join(b.get("text", "") if isinstance(b, dict) else str(b) for b in content)
    return role, content


def _tool_event(name, args):
    args = args if isinstance(args, dict) else {}
    fn = str(args.get("filename", "")).strip()
    if name == "write_todos":
        return "Planning the steps"
    if name == "read_file":
        p = str(args.get("file_path") or args.get("path") or "")
        if "/skills/" in p and p.endswith("SKILL.md"):
            return f"Reading skill: {p.split('/skills/')[1].split('/')[0]}"
        return f"Reading {os.path.basename(p) or 'a file'}"
    if name in ("ls", "glob", "grep"):
        return "Exploring the workspace"
    if name == "create_excel":
        return f"Creating Excel workbook: {fn or 'workbook.xlsx'}"
    if name == "create_word":
        return f"Writing Word document: {fn or 'document.docx'}"
    if name == "create_pdf":
        return f"Writing PDF: {fn or 'document.pdf'}"
    if name == "read_pdf":
        return f"Reading PDF{f': {fn}' if fn else ''}"
    if name == "read_spreadsheet":
        return f"Reading spreadsheet{f': {fn}' if fn else ''}"
    if name == "analyze_data":
        return f"Analyzing data{f': {fn}' if fn else ''}"
    if name == "read_document":
        return f"Reading document{f': {fn}' if fn else ''}"
    if name == "task":
        return "Delegating to a subagent"
    return f"Calling {name}"


def _events_from_messages(messages):
    out = []
    for m in messages:
        tool_calls = getattr(m, "tool_calls", None) or []
        if getattr(m, "type", "") == "ai" and tool_calls:
            for tc in tool_calls:
                if isinstance(tc, dict):
                    nm, ar = tc.get("name", "tool"), tc.get("args", {})
                else:
                    nm, ar = getattr(tc, "name", "tool"), getattr(tc, "args", {})
                out.append(("🔧", _tool_event(nm, ar)))
        elif getattr(m, "type", "") == "tool":
            out.append(("✓", f"{getattr(m, 'name', 'tool')} finished"))
    return out


def _final_answer(m):
    if getattr(m, "type", "") != "ai" or getattr(m, "tool_calls", None):
        return ""
    content = getattr(m, "content", "")
    if isinstance(content, list):
        content = "".join(b.get("text", "") if isinstance(b, dict) else str(b) for b in content)
    return content or ""


# self-heal coerced message objects back into dicts (preserving any saved steps)
def _norm_msg(m):
    r, c = _msg_parts(m)
    out = {"role": r, "content": c}
    if isinstance(m, dict) and m.get("steps"):
        out["steps"] = m["steps"]
    return out


st.session_state.da_messages = [_norm_msg(m) for m in st.session_state.da_messages]
if not st.session_state.da_messages:
    st.session_state.da_messages.append({
        "role": "assistant",
        "content": ("Hi! I'm a deep agent. Upload a file or tell me what to build — I'll plan it, "
                    "use the right skill, and drop any files I create into the panel on the right."),
    })


# =============================================================================
# SIDE PANEL (Files w/ preview, Plan, Skills)
# =============================================================================

_ICON = {".xlsx": "📊", ".csv": "📈", ".pdf": "📕", ".docx": "📄", ".txt": "📝"}
_MIME = {
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".pdf": "application/pdf", ".csv": "text/csv", ".txt": "text/plain",
}


def _render_preview(path: str):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".xlsx":
            for nm, df in _read_excel(path, sheet_name=None).items():
                st.caption(f"Sheet: {nm} — {df.shape[0]}×{df.shape[1]}")
                st.dataframe(df.head(50), use_container_width=True, height=200)
        elif ext == ".csv":
            df = pd.read_csv(path)
            st.dataframe(df.head(50), use_container_width=True, height=200)
        elif ext == ".docx":
            txt = "\n".join(p.text for p in Document(path).paragraphs)
            st.text((txt[:3000] + "…") if len(txt) > 3000 else (txt or "(empty)"))
        elif ext == ".txt":
            txt = open(path, encoding="utf-8", errors="replace").read()
            st.text((txt[:3000] + "…") if len(txt) > 3000 else txt)
        elif ext == ".pdf":
            pages = PdfReader(path).pages
            t = (pages[0].extract_text() or "") if pages else ""
            st.text((t[:2000] + "…") if len(t) > 2000 else (t or "(no extractable text)"))
        else:
            st.caption("No preview available.")
    except Exception as e:
        st.caption(f"Preview unavailable: {type(e).__name__}")


def _render_one_file(n, seen):
    path = os.path.join(WORKSPACE, n)
    ext = os.path.splitext(n)[1].lower()
    is_new = n not in seen
    size = os.path.getsize(path)
    size_str = f"{size / 1024:.1f} KB" if size < 1_048_576 else f"{size / 1_048_576:.1f} MB"
    label = f"{_ICON.get(ext, '📎')}  {n}" + ("   🟢 new" if is_new else "")
    with st.expander(label, expanded=is_new):
        st.markdown(f'<span class="file-meta">{size_str}</span>', unsafe_allow_html=True)
        _render_preview(path)
        with open(path, "rb") as fh:
            data = fh.read()
        c_dl, c_rm = st.columns(2)
        with c_dl:
            st.download_button("⬇  Download", data, file_name=n,
                               mime=_MIME.get(ext, "application/octet-stream"),
                               key=f"dl_{n}_{len(data)}", use_container_width=True)
        with c_rm:
            if st.button("🗑  Remove", key=f"rm_{n}", use_container_width=True):
                try:
                    os.remove(path)
                except OSError:
                    pass
                st.session_state.da_uploaded = [
                    x for x in st.session_state.get("da_uploaded", []) if x != n]
                st.session_state.da_seen_files = [
                    x for x in st.session_state.get("da_seen_files", []) if x != n]
                st.rerun()


def _render_files():
    files = _workspace_files()
    files = sorted(files, key=lambda n: os.path.getmtime(os.path.join(WORKSPACE, n)), reverse=True)
    seen = set(st.session_state.get("da_seen_files", []))
    uploaded_set = set(st.session_state.get("da_uploaded", []))
    generated = [n for n in files if n not in uploaded_set]
    uploaded = [n for n in files if n in uploaded_set]
    with st.container(border=True):
        st.markdown(f'<div class="panel-label">Files <span>{len(files)}</span></div>',
                    unsafe_allow_html=True)
        t_gen, t_up = st.tabs([f"Agent-generated ({len(generated)})", f"Your uploads ({len(uploaded)})"])
        with t_gen:
            if not generated:
                st.markdown('<div class="empty">Nothing generated yet — ask the agent to build something.</div>',
                            unsafe_allow_html=True)
            for n in generated:
                _render_one_file(n, seen)
        with t_up:
            if not uploaded:
                st.markdown('<div class="empty">No uploaded files — use the uploader above.</div>',
                            unsafe_allow_html=True)
            for n in uploaded:
                _render_one_file(n, seen)
    st.session_state.da_seen_files = files


def _render_plan():
    todos = st.session_state.da_todos or []
    with st.container(border=True):
        st.markdown(f'<div class="panel-label">Plan <span>{len(todos)}</span></div>',
                    unsafe_allow_html=True)
        if not todos:
            st.markdown('<div class="empty">No active plan. The agent adds steps for multi-step tasks.</div>',
                        unsafe_allow_html=True)
            return
        cls = {"completed": "d-done", "in_progress": "d-prog", "pending": "d-todo"}
        rows = []
        for t in todos:
            if not isinstance(t, dict):
                continue
            status = t.get("status", "pending")
            done = "done" if status == "completed" else ""
            rows.append(f'<div class="todo-row {done}"><span class="dot {cls.get(status, "d-todo")}"></span>'
                        f'{t.get("content", "")}</div>')
        st.markdown("".join(rows), unsafe_allow_html=True)


def _render_skills():
    with st.container(border=True):
        st.markdown('<div class="panel-label">Skills <span>click to view</span></div>',
                    unsafe_allow_html=True)
        for name, desc, body in _skill_details():
            with st.popover(f"📘  {name}", use_container_width=True):
                st.markdown(f"#### {name}")
                if desc:
                    st.caption(desc)
                st.markdown(body if body else "_No instructions found._")


def _render_debug():
    trace = st.session_state.get("da_last_trace", [])
    reads = st.session_state.get("da_last_skill_reads", [])
    if not trace:
        return
    verdict = ("Skills used: " + ", ".join(reads)) if reads else "No skill matched last run"
    with st.expander(f"🔍 Last run activity — {verdict}", expanded=False):
        lines = [f"{'→ call  ' if k == 'call' else '← result'} {n}: {d}" for k, n, d in trace]
        st.code("\n".join(lines) or "(no tool activity)", language="text")


# =============================================================================
# LAYOUT: chat (left) + side panel (right)
# =============================================================================

col_chat, col_panel = st.columns([1.85, 1], gap="large")

with col_panel:
    uploaded_files = st.file_uploader(
        "Add files (xlsx, csv, pdf, docx, txt)",
        type=["xlsx", "csv", "pdf", "docx", "txt"],
        accept_multiple_files=True,
        label_visibility="visible",
    )
    if uploaded_files:
        up = set(st.session_state.get("da_uploaded", []))
        for f in uploaded_files:
            dest = _ws_path(f.name)
            if not os.path.isfile(dest):
                with open(dest, "wb") as out:
                    out.write(f.getvalue())
            up.add(os.path.basename(dest))
        st.session_state.da_uploaded = list(up)
    _render_files()
    _render_plan()
    _render_skills()
    _render_debug()

with col_chat:
    for message in st.session_state.da_messages:
        if not message["content"]:
            continue
        with st.chat_message(message["role"]):
            if message.get("steps"):
                with st.expander(f"🔧 Steps it took ({len(message['steps'])})", expanded=False):
                    for s in message["steps"]:
                        st.markdown(f"- {s}")
            st.write(message["content"])


# =============================================================================
# USER INPUT  (docks to the bottom)
# =============================================================================

user_input = st.chat_input("Ask the deep agent to build or analyze something...")

if user_input:
    st.session_state.da_messages.append({"role": "user", "content": user_input})
    with col_chat:
        with st.chat_message("user"):
            st.write(user_input)
        with st.chat_message("assistant"):
            status = st.status("Planning and working…", expanded=True)
            response_text, steps = "", 0
            skill_reads, trace, steps_log = [], [], []
            try:
                agent_msgs = [dict(m) for m in st.session_state.da_messages]
                files = _workspace_files()
                matched = _match_skills(user_input, files)

                ctx = ["Workspace files: " + (", ".join(files) if files else "(none)")
                       + ". Use these EXACT filenames — never invent one."]
                for sk in matched:
                    body = _load_skill_body(sk)
                    if body:
                        ctx.append(f"=== Skill loaded: {sk} (follow these instructions) ===\n{body}")
                prefix = "[Context for this task]\n" + "\n\n".join(ctx) + "\n\n[User request]\n"
                if agent_msgs and agent_msgs[-1]["role"] == "user":
                    agent_msgs[-1] = {**agent_msgs[-1], "content": prefix + agent_msgs[-1]["content"]}

                for sk in matched:
                    status.write(f"📖 Loaded skill: {sk}")
                    steps_log.append(f"📖 Loaded skill: {sk}")
                _dbg("=== new turn === user:", user_input, "| files:", files, "| skills:", matched)

                for chunk in st.session_state.da_agent.stream(
                        {"messages": agent_msgs},
                        config={"recursion_limit": 100}, stream_mode="updates"):
                    for _node, update in chunk.items():
                        if not isinstance(update, dict):
                            continue
                        _dbg("node:", _node, "| keys:", list(update.keys()))
                        if update.get("todos"):
                            st.session_state.da_todos = update["todos"]
                        msgs = update.get("messages", []) or []
                        for m in msgs:
                            for tc in (getattr(m, "tool_calls", None) or []):
                                nm = tc.get("name") if isinstance(tc, dict) else getattr(tc, "name", "")
                                ar = tc.get("args", {}) if isinstance(tc, dict) else getattr(tc, "args", {})
                                _dbg("  TOOL CALL ->", nm, "|", json.dumps(ar, default=str)[:300])
                                trace.append(("call", nm, json.dumps(ar, default=str)[:300]))
                                if nm == "read_file":
                                    p = str((ar or {}).get("file_path") or (ar or {}).get("path") or "")
                                    if "/skills/" in p and p.endswith("SKILL.md"):
                                        skill_reads.append(p.split("/skills/")[1].split("/")[0])
                            if getattr(m, "type", "") == "tool":
                                prev = (m.content if isinstance(m.content, str) else str(m.content))
                                _dbg("  TOOL RESULT <-", getattr(m, "name", "?"), "|", prev[:200].replace("\n", " "))
                                trace.append(("result", getattr(m, "name", "?"), prev[:200]))
                        for icon, text in _events_from_messages(msgs):
                            steps += 1
                            status.write(f"{icon} {text}")
                            status.update(label=text)
                            steps_log.append(f"{icon} {text}")
                        for m in msgs:
                            ans = _final_answer(m)
                            if ans:
                                response_text = ans

                applied = sorted(set(skill_reads) | set(matched))
                st.session_state.da_last_trace = trace
                st.session_state.da_last_skill_reads = applied
                status.update(label=f"Done · {steps} step(s)"
                              + (f" · skills: {', '.join(applied)}" if applied else ""),
                              state="complete", expanded=False)
                response_text = response_text or "(the agent finished without a textual reply)"
                st.write(response_text)
                st.session_state.da_messages.append(
                    {"role": "assistant", "content": response_text, "steps": steps_log})
            except Exception as e:
                status.update(label="Error", state="error")
                _dbg("ERROR:", repr(e))
                st.error(f"Error: {type(e).__name__}: {e}")
                st.session_state.da_messages.append(
                    {"role": "assistant", "content": f"Error: {type(e).__name__}: {e}"})
    st.rerun()
